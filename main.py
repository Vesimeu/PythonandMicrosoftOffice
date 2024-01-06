import sys
import os
from PyQt5.QtWidgets import QApplication, QMainWindow, QPushButton, QVBoxLayout, QFileDialog, QWidget, QMessageBox, \
    QListWidget
import pandas as pd
from docx import Document


class ExcelToWordApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.initUI()

    def initUI(self):
        self.setWindowTitle('Конвертер Excel в Word')
        self.setGeometry(100, 100, 800, 600)
        layout = QVBoxLayout()

        self.list_widget = QListWidget(self)
        layout.addWidget(self.list_widget)

        btn_select_excel = QPushButton('Выберите файл Excel', self)
        btn_select_excel.clicked.connect(self.openFileNameDialog)
        layout.addWidget(btn_select_excel)

        self.btn_generate_word = QPushButton('Генерировать Word документ', self)
        self.btn_generate_word.clicked.connect(self.generateWordDocument)
        self.btn_generate_word.setEnabled(False)  # Кнопка неактивна пока не выбран Excel файл
        layout.addWidget(self.btn_generate_word)

        self.central_widget = QWidget()
        self.central_widget.setLayout(layout)
        self.setCentralWidget(self.central_widget)

    def openFileNameDialog(self):
        options = QFileDialog.Options()
        fileName, _ = QFileDialog.getOpenFileName(self, "Выберите файл Excel", "", "Excel Files (*.xlsx)",
                                                  options=options)
        if fileName:
            self.excel_file = fileName
            self.loadExcelData(fileName)

    def loadExcelData(self, excel_file):
        self.df = pd.read_excel(excel_file)  # Считывание Excel
        self.displayData()

    def displayData(self):
        # Отображение данных в виджете списка
        self.list_widget.clear()
        for index, row in self.df.iterrows():
            self.list_widget.addItem(f'Строка {index + 1}: {row.to_dict()}')
        self.btn_generate_word.setEnabled(True)  # Активация кнопки после загрузки данных

    # Функция для генерации документа
    def generateWordDocument(self):
        selected_item = self.list_widget.currentItem()
        if not selected_item:
            QMessageBox.warning(self, 'Ошибка', 'Выберите строку данных для генерации документа.')
            return

        current_row = self.list_widget.currentRow()
        record = self.df.iloc[current_row].to_dict()

        # Загрузка шаблона Word
        word_template = "output.docx"  # Укажите путь к вашему шаблону Word
        document = Document(word_template)

        # Заполнение шаблона данными
        for paragraph in document.paragraphs:
            # Замена плейсхолдеров в тексте параграфа
            for key, value in record.items():
                if f'{{{key}}}' in paragraph.text:
                    paragraph.text = paragraph.text.replace(f'{{{key}}}', str(value))

        # Также проверяем таблицы в документе
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in record.items():
                        if f'{{{key}}}' in cell.text:
                            cell.text = cell.text.replace(f'{{{key}}}', str(value))

        # Сохранение нового документа Word
        new_file_path = os.path.join(os.getcwd(), f'Справка_{record["ФИО"]}.docx')
        document.save(new_file_path)
        QMessageBox.information(self, 'Готово', f'Документ сохранен: {new_file_path}')


def main():
    app = QApplication(sys.argv)
    main = ExcelToWordApp()
    main.show()
    sys.exit(app.exec_())


if __name__ == '__main__':
    main()
